from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib.auth import authenticate, login, logout
from django.contrib.admin.views.decorators import staff_member_required
from django.contrib.auth.mixins import LoginRequiredMixin, UserPassesTestMixin
from django.contrib import messages
from django.http import HttpResponse, JsonResponse
from django.template.loader import get_template, render_to_string
from django.views.decorators.http import require_POST
from django.utils.html import strip_tags
from django.contrib.sites.shortcuts import get_current_site
from django.views.generic import ListView, DetailView, CreateView, UpdateView, DeleteView
from django.db.models import Count, Q
from django.urls import reverse_lazy
from django.core.mail import EmailMultiAlternatives
from django.conf import settings
from django_filters.views import FilterView
from django import forms
import docx
import pdfplumber
import io
import logging
import yake

logger = logging.getLogger('journalapp')

from .models import (
    Article, Department, Profile, ArticleCategory, Review, SiteSettings,
    HeroSlide, Journal, Issue, ArticleLog, Rubric, Submission,
    Assignment, SubmissionMessage
)
from .permissions import (
    journals_for, can_manage_any_journal, WORKFLOW_ROLES,
)
from .utils import get_from_email
from .forms import (
    UserRegisterForm, ProfileUpdateForm, UserUpdateForm, ArticleForm,
    ReviewForm, CommentForm, DepartmentForm, SiteSettingsForm,
    HeroSlideForm, AssignReviewerForm
)
from .filters import ArticleFilter


@login_required
@require_POST
def parse_document(request):
    if 'document' not in request.FILES:
        return JsonResponse({'error': 'No document uploaded'}, status=400)

    file = request.FILES['document']
    file_name = file.name.lower()
    content = ""

    try:
        if file_name.endswith('.docx'):
            doc = docx.Document(io.BytesIO(file.read()))
            full_text = [para.text for para in doc.paragraphs]
            content = '\n'.join(full_text)
        elif file_name.endswith('.pdf'):
            with pdfplumber.open(io.BytesIO(file.read())) as pdf:
                full_text = [page.extract_text() for page in pdf.pages if page.extract_text()]
                content = '\n'.join(full_text)
        else:
            return JsonResponse({'error': 'Unsupported file type. Please upload a .docx or .pdf file.'}, status=400)

        lines = [line.strip() for line in content.split('\n') if line.strip()]
        title = lines[0] if lines else 'Could not determine title'
        
        abstract = ""
        body = content
        # A simple logic to separate abstract
        try:
            abstract_index = [i for i, line in enumerate(lines) if 'abstract' in line.lower()][0]
            # Assuming abstract is the content after the 'abstract' keyword line
            # and body starts after a blank line or next heading.
            abstract_lines = []
            for line in lines[abstract_index+1:]:
                if not line.strip(): # break on blank line
                    break
                abstract_lines.append(line)
            abstract = "\n".join(abstract_lines)
        except IndexError:
            abstract = "" # No abstract found

        kw_extractor = yake.KeywordExtractor(lan="en", n=1, dedupLim=0.9, top=10)
        keywords = kw_extractor.extract_keywords(content)
        keyword_list = [kw for kw, score in keywords]

        return JsonResponse({
            'title': title,
            'abstract': abstract,
            'content': body,
            'keywords': ", ".join(keyword_list)
        })

    except Exception as e:
        return JsonResponse({'error': f'An error occurred while parsing the file: {str(e)}'}, status=500)


def home(request):
    """Landing page. Leads with the journals themselves.

    This used to list Departments, which hid the fact that one department can
    run several journals — a visitor could not tell JJEL from JOJWOL. Journals
    are now the first thing on the page, each linking to its own site.
    """
    # order_by is explicit because Django drops Meta.ordering from a query that
    # groups (which annotate() makes it do), so the cards would come back in
    # whatever order the database happened to return.
    journals = Journal.objects.filter(is_active=True).annotate(
        article_count=Count('articles', filter=Q(articles__status='published'), distinct=True),
        issue_count=Count('issues', filter=Q(issues__is_published=True), distinct=True),
    ).order_by('order', 'name')
    latest_articles = (
        Article.objects.filter(status='published')
        .select_related('journal', 'issue')
        .order_by('-published_at')[:6]
    )
    hero_slides = HeroSlide.objects.filter(is_active=True).order_by('order')
    featured_issues = (
        Issue.objects.filter(featured=True, is_published=True, journal__is_active=True)
        .select_related('journal')[:3]
    )

    context = {
        'journals': journals,
        'latest_articles': latest_articles,
        'hero_slides': hero_slides,
        'featured_issues': featured_issues,
    }
    return render(request, 'journalapp/home.html', context)


################ Authentication Views #######################
def register(request):
    if request.method == 'POST':
        form = UserRegisterForm(request.POST)
        if form.is_valid():
            user = form.save()
            # Send welcome email
            send_welcome_email(request, user)
            messages.success(request, f'Your account has been created! You can now log in.')
            return redirect('login')
    else:
        form = UserRegisterForm()
    
    context = {
        'form': form,
        'departments': Department.objects.all()
    }
    return render(request, 'journalapp/register.html', context)


def send_welcome_email(request, user):
    current_site = get_current_site(request)
    site_url = f"{'https' if request.is_secure() else 'http'}://{current_site.domain}"
    
    context = {
        'user': user,
        'site_url': site_url,
    }
    
    # Render email templates
    html_content = render_to_string('journalapp/welcome_email.html', context)
    text_content = render_to_string('journalapp/welcome_email.txt', context)
    
    # Create email — the From address MUST be a verified sender in your SMTP
    # provider (e.g. Brevo), otherwise the send is rejected. Driven by
    # DEFAULT_FROM_EMAIL in settings/.env.
    subject = "Welcome to University Journal"
    from_email = get_from_email()
    to_email = user.email

    # Send email
    email = EmailMultiAlternatives(subject, text_content, from_email, [to_email])
    email.attach_alternative(html_content, "text/html")

    try:
        email.send()
        logger.info("Welcome email sent to %s", to_email)
    except Exception as e:
        # Log the error but don't prevent user registration.
        logger.error("Error sending welcome email to %s: %s", to_email, e, exc_info=True)


def login_view(request):
    if request.method == 'POST':
        email = request.POST.get('email')
        password = request.POST.get('password')
        remember = request.POST.get('remember')
        
        # Try to authenticate with email directly
        user = authenticate(request, email=email, password=password)
        
        # If that fails, try with username parameter (Django's default)
        if user is None:
            user = authenticate(request, username=email, password=password)
        
        if user is not None:
            login(request, user)
            # Get the next parameter or default to dashboard
            next_url = request.POST.get('next') or request.GET.get('next')
            if next_url:
                return redirect(next_url)
            else:
                return redirect('dashboard')
        else:
            messages.error(request, 'Invalid username or password.')
    
    return render(request, 'journalapp/login.html')


def logout_view(request):
    logout(request)
    messages.success(request, "You have been successfully logged out.")
    return redirect('home')

@login_required
def profile(request):
    if request.method == 'POST':
        u_form = UserUpdateForm(request.POST, instance=request.user)
        p_form = ProfileUpdateForm(request.POST, request.FILES, instance=request.user.profile)
        
        if u_form.is_valid() and p_form.is_valid():
            u_form.save()
            p_form.save()
            messages.success(request, f'Your profile has been updated!')
            return redirect('profile')
    else:
        u_form = UserUpdateForm(instance=request.user)
        p_form = ProfileUpdateForm(instance=request.user.profile)
    
    context = {
        'u_form': u_form,
        'p_form': p_form
    }
    return render(request, 'journalapp/profile.html', context)


################ End of Authentication Views #######################

@login_required
def dashboard(request):
    if request.user.is_staff or can_manage_any_journal(request.user, roles=WORKFLOW_ROLES):
        return redirect('admin_dashboard')
    elif getattr(request.user, 'profile', None) and request.user.profile.is_reviewer:
        return redirect('reviewer_dashboard')
    elif getattr(request.user, 'profile', None) and request.user.profile.is_editor:
        return redirect('editor_dashboard')
    else:
        return redirect('author_dashboard')

@login_required
def author_dashboard(request):
    # Get user's submissions (new workflow)
    user_submissions = Submission.objects.filter(
        author=request.user
    ).select_related('journal').order_by('-submitted_at')

    # Get user's articles (legacy - for backward compatibility)
    user_articles = Article.objects.filter(author=request.user).order_by('-created_at')

    journals = Journal.objects.all()

    # Recent submissions (last 5)
    recent_submissions = user_submissions[:5]

    # Submissions requiring action (revision requested) — actionable, not a
    # status browser, so this stays even though the stats tiles were removed.
    action_required = user_submissions.filter(status='revision_requested')

    # Submissions with feedback available (completed assignments)
    feedback_available = user_submissions.filter(
        assignments__status='completed'
    ).distinct().prefetch_related(
        'assignments__assigned_to'
    ).order_by('-updated_at')[:5]

    context = {
        'user_submissions': user_submissions,
        'recent_submissions': recent_submissions,
        'action_required': action_required,
        'feedback_available': feedback_available,
        'user_articles': user_articles,
        'journals': journals,
        # Legacy stat still used by the "My published articles" section.
        'published_count': user_articles.filter(status='published').count(),
    }
    return render(request, 'journalapp/author_dashboard.html', context)

@login_required
def editor_dashboard(request):
    if not (getattr(request.user, 'profile', None) and request.user.profile.is_editor):
        return redirect('dashboard')

    # Get assignments for this editor
    active_assignments = Assignment.objects.filter(
        assigned_to=request.user,
        role='editor',
        status='active'
    ).select_related('submission', 'submission__author', 'submission__journal').order_by('-assigned_at')

    # Get all completed assignments for stats (before slicing)
    all_completed = Assignment.objects.filter(
        assigned_to=request.user,
        role='editor',
        status='completed'
    )

    # Get recent 10 completed for display
    completed_assignments = all_completed.select_related('submission').order_by('-completed_at')[:10]

    # Calculate stats
    assignment_stats = {
        'active': active_assignments.count(),
        'completed': all_completed.count(),  # Use all_completed, not sliced
        'total': Assignment.objects.filter(assigned_to=request.user, role='editor').count(),
    }

    # Count by recommendation for completed (use all_completed for accurate stats)
    recommendations = all_completed.values_list('recommendation', flat=True)
    assignment_stats['accepted'] = sum(1 for r in recommendations if r == 'accept')
    assignment_stats['rejected'] = sum(1 for r in recommendations if r == 'reject')
    assignment_stats['revisions'] = sum(1 for r in recommendations if r == 'revisions')

    # Get unread messages count
    unread_messages = SubmissionMessage.objects.filter(
        recipient=request.user,
        is_read=False
    ).count()

    # Get recent messages
    recent_messages = SubmissionMessage.objects.filter(
        Q(recipient=request.user) | Q(sender=request.user)
    ).select_related('submission', 'sender').order_by('-created_at')[:5]

    context = {
        'active_assignments': active_assignments,
        'completed_assignments': completed_assignments,
        'assignment_stats': assignment_stats,
        'unread_messages': unread_messages,
        'recent_messages': recent_messages,
    }
    return render(request, 'journalapp/editor_dashboard.html', context)

@login_required
def reviewer_dashboard(request):
    if not (getattr(request.user, 'profile', None) and request.user.profile.is_reviewer):
        return redirect('dashboard')

    # Get assignments for this reviewer
    active_assignments = Assignment.objects.filter(
        assigned_to=request.user,
        role='reviewer',
        status='active'
    ).select_related('submission', 'submission__author', 'submission__journal').order_by('-assigned_at')

    # Get all completed assignments for stats (before slicing)
    all_completed = Assignment.objects.filter(
        assigned_to=request.user,
        role='reviewer',
        status='completed'
    )

    # Get recent 10 completed for display
    completed_assignments = all_completed.select_related('submission').order_by('-completed_at')[:10]

    # Calculate stats
    assignment_stats = {
        'active': active_assignments.count(),
        'completed': all_completed.count(),  # Use all_completed, not sliced
        'total': Assignment.objects.filter(assigned_to=request.user, role='reviewer').count(),
    }

    # Count by recommendation for completed (use all_completed for accurate stats)
    recommendations = all_completed.values_list('recommendation', flat=True)
    assignment_stats['accepted'] = sum(1 for r in recommendations if r == 'accept')
    assignment_stats['rejected'] = sum(1 for r in recommendations if r == 'reject')
    assignment_stats['revisions'] = sum(1 for r in recommendations if r == 'revisions')

    # Get unread messages count
    unread_messages = SubmissionMessage.objects.filter(
        recipient=request.user,
        is_read=False
    ).count()

    # Get recent messages
    recent_messages = SubmissionMessage.objects.filter(
        Q(recipient=request.user) | Q(sender=request.user)
    ).select_related('submission', 'sender').order_by('-created_at')[:5]

    context = {
        'active_assignments': active_assignments,
        'completed_assignments': completed_assignments,
        'assignment_stats': assignment_stats,
        'unread_messages': unread_messages,
        'recent_messages': recent_messages,
    }
    return render(request, 'journalapp/reviewer_dashboard.html', context)

@login_required
def admin_dashboard(request):
    # Scope everything on this page to the journals the user may act on. Site
    # staff get every journal back from journals_for(), so their view is
    # unchanged; a journal chief editor sees only their own journal's figures.
    permitted_journals = journals_for(request.user, roles=WORKFLOW_ROLES)
    if not permitted_journals.exists():
        return redirect('dashboard')

    all_submissions = (
        Submission.objects
        .filter(journal__in=permitted_journals)
        .select_related('author', 'journal')
    )

    # Submission stats by status
    submission_stats = {
        'total': all_submissions.count(),
        'pending': all_submissions.filter(status='pending').count(),
        'in_review': all_submissions.filter(status='in_review').count(),
        'with_editor': all_submissions.filter(status='with_editor').count(),
        'revision_requested': all_submissions.filter(status='revision_requested').count(),
        'revised': all_submissions.filter(status='revised').count(),
        'approved': all_submissions.filter(status='approved').count(),
        'published': all_submissions.filter(status='published').count(),
        'rejected': all_submissions.filter(status='rejected').count(),
    }

    # Submissions needing attention
    pending_submissions = all_submissions.filter(status='pending').order_by('-submitted_at')[:5]
    revised_submissions = all_submissions.filter(status='revised').order_by('-updated_at')[:5]
    approved_submissions = all_submissions.filter(status='approved').order_by('-updated_at')[:5]

    # Completed reviews - submissions where reviewer completed but admin hasn't decided
    # These are submissions in_review or with_editor status that have at least one completed assignment
    from .models import Assignment
    completed_review_submissions = all_submissions.filter(
        status__in=['in_review', 'with_editor'],
        assignments__status='completed'
    ).distinct().order_by('-updated_at')[:5]

    # Prefetch related assignment data for displaying reviewer feedback
    completed_review_submissions = completed_review_submissions.prefetch_related(
        'assignments__assigned_to'
    )

    # Recent activity - all recent submissions
    recent_submissions = all_submissions.order_by('-updated_at')[:10]

    # Get reviewers and editors for quick reference
    from .models import Profile, GuestReviewer
    reviewers = Profile.objects.filter(is_reviewer=True).select_related('user').order_by('user__first_name', 'user__last_name')
    editors = Profile.objects.filter(is_editor=True).select_related('user').order_by('user__first_name', 'user__last_name')
    guest_reviewers = GuestReviewer.objects.filter(is_active=True).order_by('first_name', 'last_name')

    # Unread messages for admin
    unread_messages = SubmissionMessage.objects.filter(
        recipient=request.user,
        is_read=False
    ).count()

    # Journal stats
    journals = permitted_journals
    journal_stats = []
    for journal in journals:
        journal_stats.append({
            'journal': journal,
            'submissions': journal.submissions.count(),
            'pending': journal.submissions.filter(status='pending').count(),
        })

    # Add completed reviews count to stats
    submission_stats['completed_reviews'] = all_submissions.filter(
        status__in=['in_review', 'with_editor'],
        assignments__status='completed'
    ).distinct().count()

    context = {
        'submission_stats': submission_stats,
        'pending_submissions': pending_submissions,
        'revised_submissions': revised_submissions,
        'approved_submissions': approved_submissions,
        'completed_review_submissions': completed_review_submissions,
        'recent_submissions': recent_submissions,
        'reviewers': reviewers,
        'editors': editors,
        'guest_reviewers': guest_reviewers,
        'unread_messages': unread_messages,
        'journal_stats': journal_stats,
    }
    return render(request, 'journalapp/admin_dashboard.html', context)

@login_required
def assign_reviewer(request, article_id):
    if not (getattr(request.user, 'profile', None) and request.user.profile.is_editor):
        messages.error(request, "You don't have permission to perform this action.")
        return redirect('dashboard')

    article = get_object_or_404(Article, pk=article_id)
    if request.method == 'POST':
        form = AssignReviewerForm(request.POST)
        if form.is_valid():
            reviewer = form.cleaned_data['reviewer']
            if Review.objects.filter(article=article, reviewer=reviewer).exists():
                messages.warning(request, f"{reviewer.get_full_name()} is already assigned to review this article.")
            else:
                Review.objects.create(article=article, reviewer=reviewer)
                article.status = 'under_review'
                article.save()
                ArticleLog.objects.create(article=article, user=request.user, action=f"Assigned reviewer: {reviewer.get_full_name()}")
                messages.success(request, f"Assigned {reviewer.get_full_name()} to review '{article.title}'.")
            return redirect('editor_dashboard')
    else:
        form = AssignReviewerForm()

    context = {'form': form, 'article': article}
    return render(request, 'journalapp/assign_reviewer.html', context)

@login_required
def upload_revised_document(request, pk):
    article = get_object_or_404(Article, pk=pk)

    # Check if user is editor or admin
    if not request.user.profile.is_editor and not request.user.is_staff:
        messages.error(request, 'You do not have permission to upload revised documents.')
        return redirect('dashboard')

    if request.method == 'POST':
        if 'revised_document' in request.FILES:
            article.revised_document = request.FILES['revised_document']
            article.revision_notes = request.POST.get('revision_notes', '')
            article.status = 'revision_required'  # Update status
            article.save()

            # Notify the author
            send_revision_notification(request, article)

            messages.success(request, 'Revised document uploaded successfully.')
        else:
            messages.error(request, 'No document was uploaded.')

    return redirect('dashboard')


def send_revision_notification(request, article):
    current_site = get_current_site(request)
    site_url = f"{'https' if request.is_secure() else 'http'}://{current_site.domain}"
    dashboard_url = f"{site_url}/dashboard/"

    context = {
        'article': article,
        'user': article.author,
        'site_url': site_url,
        'dashboard_url': dashboard_url,
        'revision_notes': article.revision_notes,
    }

    # Render email templates
    html_content = render_to_string('journalapp/revision_notification.html', context)
    text_content = strip_tags(html_content)

    # Create email
    subject = f"Revision Required: {article.title}"
    from_email = get_from_email()
    to_email = article.author.email

    # Send email
    email = EmailMultiAlternatives(subject, text_content, from_email, [to_email])
    email.attach_alternative(html_content, "text/html")

    try:
        email.send()
        logger.info("Revision notification email sent to %s", to_email)
    except Exception as e:
        logger.error("Error sending revision notification email to %s: %s", to_email, e, exc_info=True)


class ArticleListView(FilterView):
    model = Article
    template_name = 'journalapp/article_list.html'
    context_object_name = 'articles'
    filterset_class = ArticleFilter
    paginate_by = 10

    def get_queryset(self):
        base_qs = Article.objects.filter(status='published').order_by('-published_at')
        # Apply django-filter on department/category/title (via ArticleFilter)
        self.filterset = self.filterset_class(self.request.GET, queryset=base_qs)
        queryset = self.filterset.qs
        # Simple search support via ?q=
        q = self.request.GET.get('q')
        if q:
            queryset = queryset.filter(title__icontains=q)
        # Optional date range filtering
        published_after = self.request.GET.get('published_at_after')
        published_before = self.request.GET.get('published_at_before')
        if published_after:
            queryset = queryset.filter(published_at__date__gte=published_after)
        if published_before:
            queryset = queryset.filter(published_at__date__lte=published_before)
        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['journals'] = Journal.objects.filter(is_active=True)
        slug = self.request.GET.get('journal')
        context['selected_journal'] = (
            Journal.objects.filter(slug=slug).first() if slug else None
        )
        return context

class ArticleDetailView(DetailView):
    model = Article
    template_name = 'journalapp/article_detail.html'
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['comment_form'] = CommentForm()
        context['logs'] = self.get_object().logs.all()
        return context

class ArticleCreateView(LoginRequiredMixin, CreateView):
    model = Article
    form_class = ArticleForm
    template_name = 'journalapp/article_form.html'

    def get_initial(self):
        initial = super().get_initial()
        journal = get_object_or_404(Journal, pk=self.kwargs['journal_id'])
        initial['journal'] = journal
        return initial

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        journal = get_object_or_404(Journal, pk=self.kwargs['journal_id'])
        context['journal'] = journal
        context['rubrics'] = Rubric.objects.filter(journal=journal)
        return context

    def get_form(self, form_class=None):
        form = super().get_form(form_class)
        journal = get_object_or_404(Journal, pk=self.kwargs['journal_id'])
        form.fields['category'].queryset = ArticleCategory.objects.filter(journal=journal)
        form.fields['journal'].widget = forms.HiddenInput()
        return form

    def form_valid(self, form):
        form.instance.author = self.request.user
        response = super().form_valid(form)
        ArticleLog.objects.create(article=self.object, user=self.request.user, action="Article created as draft.")
        messages.success(self.request, 'Article created successfully!')
        return response

    def get_success_url(self):
        return reverse_lazy('dashboard')

class ArticleUpdateView(LoginRequiredMixin, UserPassesTestMixin, UpdateView):
    model = Article
    form_class = ArticleForm
    template_name = 'journalapp/article_form.html'
    success_url = reverse_lazy('dashboard')

    def form_valid(self, form):
        messages.success(self.request, 'Article updated successfully!')
        return super().form_valid(form)

    def test_func(self):
        article = self.get_object()
        return self.request.user == article.author

class ArticleDeleteView(LoginRequiredMixin, UserPassesTestMixin, DeleteView):
    model = Article
    template_name = 'journalapp/article_confirm_delete.html'
    success_url = reverse_lazy('dashboard')

    def test_func(self):
        article = self.get_object()
        return self.request.user == article.author

@login_required
def submit_article(request, pk):
    article = get_object_or_404(Article, pk=pk, author=request.user)
    if article.status == 'draft':
        article.status = 'submitted'
        article.save()
        ArticleLog.objects.create(article=article, user=request.user, action="Article submitted for review.")
        messages.success(request, 'Article submitted successfully!')
    return redirect('dashboard')

@login_required
def review_article(request, pk):
    article = get_object_or_404(Article, pk=pk)
    if not (getattr(request.user, 'profile', None) and request.user.profile.is_reviewer):
        messages.error(request, 'You do not have permission to review journals.')
        return redirect('dashboard')
    
    review, created = Review.objects.get_or_create(article=article, reviewer=request.user)

    if request.method == 'POST':
        form = ReviewForm(request.POST, instance=review)
        if form.is_valid():
            form.save()
            decision = form.cleaned_data.get('decision')
            if decision == 'accept':
                article.status = 'accepted'
            elif decision in ['minor_revision', 'major_revision']:
                article.status = 'revision_required'
            elif decision == 'reject':
                article.status = 'rejected'
            article.save()
            ArticleLog.objects.create(article=article, user=request.user, action=f"Review submitted with decision: {decision.replace('_', ' ').title()}", notes=review.comments)
            messages.success(request, 'Review submitted successfully!')
            return redirect('dashboard')
    else:
        form = ReviewForm(instance=review)
    
    context = {'form': form, 'article': article}
    return render(request, 'journalapp/review_form.html', context)

@login_required
def publish_article(request, pk):
    article = get_object_or_404(Article, pk=pk)
    if not (getattr(request.user, 'profile', None) and request.user.profile.is_editor):
        messages.error(request, 'You do not have permission to publish articles.')
        return redirect('dashboard')
    
    if article.status == 'accepted':
        article.publish()
        ArticleLog.objects.create(article=article, user=request.user, action="Article published.")
        messages.success(request, 'Article published successfully!')
    return redirect('dashboard')


############## HERO SLIDER #############
@staff_member_required
def site_settings(request):
    settings, created = SiteSettings.objects.get_or_create(pk=1)

    if request.method == 'POST':
        form = SiteSettingsForm(request.POST, request.FILES, instance=settings)
        if form.is_valid():
            form.save()
            messages.success(request, 'Site settings updated successfully!')
            return redirect('site_settings')
    else:
        form = SiteSettingsForm(instance=settings)

    context = {
        'form': form,
        'title': 'Site Settings'
    }
    return render(request, 'journalapp/site_settings.html', context)

@staff_member_required
def hero_slides(request):
    slides = HeroSlide.objects.all().order_by('order')

    context = {
        'slides': slides,
        'title': 'Hero Slides'
    }
    return render(request, 'journalapp/hero_slides.html', context)

@staff_member_required
def hero_slide_create(request):
    if request.method == 'POST':
        form = HeroSlideForm(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            messages.success(request, 'Hero slide created successfully!')
            return redirect('hero_slides')
    else:
        form = HeroSlideForm()

    context = {
        'form': form,
        'title': 'Create Hero Slide'
    }
    return render(request, 'journalapp/hero_slide_form.html', context)

@staff_member_required
def hero_slide_update(request, pk):
    slide = get_object_or_404(HeroSlide, pk=pk)

    if request.method == 'POST':
        form = HeroSlideForm(request.POST, request.FILES, instance=slide)
        if form.is_valid():
            form.save()
            messages.success(request, 'Hero slide updated successfully!')
            return redirect('hero_slides')
    else:
        form = HeroSlideForm(instance=slide)

    context = {
        'form': form,
        'title': 'Update Hero Slide'
    }
    return render(request, 'journalapp/hero_slide_form.html', context)

@staff_member_required
def hero_slide_delete(request, pk):
    slide = get_object_or_404(HeroSlide, pk=pk)

    if request.method == 'POST':
        slide.delete()
        messages.success(request, 'Hero slide deleted successfully!')
        return redirect('hero_slides')

    context = {
        'slide': slide,
        'title': 'Delete Hero Slide'
    }
    return render(request, 'journalapp/hero_slide_confirm_delete.html', context)



# class AllJournalsListView(ListView):
#     model = Article
#     template_name = 'journalapp/all_journals.html'
#     context_object_name = 'journals'
#     paginate_by = 10
    
#     def get_queryset(self):
#         queryset = Article.objects.filter(status='published')
        
#         # Filter by department
#         department_id = self.request.GET.get('department')
#         if department_id:
#             queryset = queryset.filter(department_id=department_id)
        
#         # Search functionality
#         search_query = self.request.GET.get('search')
#         if search_query:
#             queryset = queryset.filter(
#                 Q(title__icontains=search_query) |
#                 Q(abstract__icontains=search_query) |
#                 Q(content__icontains=search_query)
#             )
        
#         # Sorting
#         sort_by = self.request.GET.get('sort', 'newest')
#         if sort_by == 'newest':
#             queryset = queryset.order_by('-published_at')
#         elif sort_by == 'oldest':
#             queryset = queryset.order_by('published_at')
#         elif sort_by == 'title':
#             queryset = queryset.order_by('title')
#         else:
#             queryset = queryset.order_by('-published_at')
        
#         return queryset
    
#     def get_context_data(self, **kwargs):
#         context = super().get_context_data(**kwargs)
#         context['departments'] = Department.objects.all()
#         return context


############# PDF Export ###############
def article_pdf(request, pk):
    from io import BytesIO
    from xhtml2pdf import pisa

    article = get_object_or_404(Article, pk=pk)
    template = get_template('journalapp/article_pdf.html')
    context = {
        'article': article,
        'request': request,
    }
    html = template.render(context)

    # Create a PDF
    result = BytesIO()
    pdf = pisa.pisaDocument(BytesIO(html.encode("UTF-8")), result)

    if not pdf.err:
        response = HttpResponse(result.getvalue(), content_type='application/pdf')
        filename = f"{article.title.replace(' ', '_')}.pdf"
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response

    return HttpResponse("Error generating PDF", status=400)

class ArticleSearchView(ListView):
    model = Article
    template_name = 'journalapp/article_search.html'
    context_object_name = 'articles'
    paginate_by = 10

    def get_queryset(self):
        queryset = Article.objects.filter(status='published')

        # Search functionality
        search_query = self.request.GET.get('q', '')
        if search_query:
            queryset = queryset.filter(
                Q(title__icontains=search_query) |
                Q(abstract__icontains=search_query) |
                Q(content__icontains=search_query) |
                Q(author__first_name__icontains=search_query) |
                Q(author__last_name__icontains=search_query) |
                Q(keywords__icontains=search_query)
            )

        # Filter by journal. Was department, which could not separate two
        # journals of the same department.
        journal_slug = self.request.GET.get('journal')
        if journal_slug:
            queryset = queryset.filter(journal__slug=journal_slug)

        # Filter by category
        category_id = self.request.GET.get('category')
        if category_id:
            queryset = queryset.filter(category_id=category_id)

        # Sorting
        sort_by = self.request.GET.get('sort', 'newest')
        if sort_by == 'newest':
            queryset = queryset.order_by('-published_at')
        elif sort_by == 'oldest':
            queryset = queryset.order_by('published_at')
        elif sort_by == 'title':
            queryset = queryset.order_by('title')
        elif sort_by == 'author':
            queryset = queryset.order_by('author__last_name', 'author__first_name')
        else:
            queryset = queryset.order_by('-published_at')

        return queryset.distinct()

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['journals'] = Journal.objects.filter(is_active=True)
        context['categories'] = ArticleCategory.objects.all()
        context['search_query'] = self.request.GET.get('q', '')
        context['selected_journal'] = self.request.GET.get('journal', '')
        context['selected_category'] = self.request.GET.get('category', '')
        context['selected_sort'] = self.request.GET.get('sort', 'newest')
        return context

